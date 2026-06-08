from __future__ import annotations

import json
import re
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import AppConfig
from .models import ApplicationPacket, JobListing, JobScore, Profile


def extract_resume_text(profile: Profile, root_dir: str | Path) -> str:
    resume_path = profile.resume_base_path or profile.resume_path
    if not resume_path:
        return profile.summary
    candidate = Path(resume_path)
    if not candidate.is_absolute():
        candidate = Path(root_dir) / candidate
    if not candidate.exists():
        return profile.summary
    if candidate.suffix.lower() == ".docx":
        document = Document(candidate)
        text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        return text or profile.summary
    return candidate.read_text(encoding="utf-8")


def extract_form_questions(job: JobListing) -> list[str]:
    text = _strip_html(job.description)
    questions: list[str] = []
    for line in text.splitlines():
        cleaned = " ".join(line.split()).strip(" -")
        lower = cleaned.lower()
        if not cleaned:
            continue
        if any(
            phrase in lower
            for phrase in (
                "years of experience",
                "authorized to work",
                "visa sponsorship",
                "notice period",
                "current ctc",
                "expected ctc",
                "salary expectation",
                "linkedin",
                "portfolio",
                "github",
                "cover letter",
                "why do you want",
                "start date",
            )
        ):
            questions.append(cleaned)
    return questions[:12]


def default_form_answers(profile: Profile) -> dict[str, str]:
    answers: dict[str, str] = {}
    links = []
    for keyword in profile.keywords:
        if keyword.startswith("http://") or keyword.startswith("https://"):
            links.append(keyword)
    if profile.location:
        answers["Current location"] = profile.location
    if profile.seniority:
        answers["Seniority"] = profile.seniority
    if profile.desired_salary_min > 0:
        answers["Minimum salary expectation"] = f"{profile.desired_salary_currency} {profile.desired_salary_min:,.0f}"
    if links:
        answers["Relevant links"] = ", ".join(links[:3])
    return answers


def compile_docx_resume(packet: ApplicationPacket, output_path: Path, base_docx_path: Path | None = None) -> Path:
    """
    Build a tailored resume .docx file.
    If a base_docx_path exists, clone it and prepend the tailored summary.
    Otherwise generate a clean professional template from scratch.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if base_docx_path and base_docx_path.exists():
        doc = Document(base_docx_path)
        # Insert tailored summary at the top
        tailored_intro = doc.add_paragraph()
        tailored_intro.paragraph_format.space_after = Pt(8)
        run = tailored_intro.add_run("Tailored Summary")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(packet.tailored_resume_text or packet.score.explanation())
        doc.paragraphs[0]._element.addprevious(tailored_intro._element)
    else:
        doc = Document()

        # Header — Name
        h = doc.add_heading(packet.job.score.job.company and packet.score.job.title or "Resume", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Contact / meta line
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"{packet.score.job.location}  ·  {packet.score.job.company}").font.size = Pt(10)

        doc.add_paragraph()  # spacer

        # Tailored Summary section
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(packet.tailored_resume_text or packet.score.explanation())

        # Cover Letter section
        doc.add_heading("Cover Letter", level=2)
        doc.add_paragraph(packet.cover_letter)

        # Form Answers section
        if packet.form_answers:
            doc.add_heading("Application Answers", level=2)
            for key, value in packet.form_answers.items():
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{key}: ").bold = True
                p.add_run(value)

        # Resume Notes section
        if packet.resume_notes:
            doc.add_heading("Resume Notes", level=2)
            for note in packet.resume_notes:
                doc.add_paragraph(note, style="List Bullet")

    doc.save(output_path)
    return output_path


def compile_pdf_resume(packet: ApplicationPacket, output_path: Path) -> Path:
    """
    Render the application packet as HTML and use Playwright's headless print-to-PDF.
    Falls back to writing a plain .txt file if Playwright is not available.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    html = _build_resume_html(packet)
    html_tmp = output_path.with_suffix(".html")
    html_tmp.write_text(html, encoding="utf-8")

    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(html_tmp.as_uri())
            page.wait_for_timeout(500)
            page.pdf(
                path=str(output_path),
                format="A4",
                margin={"top": "20mm", "bottom": "20mm", "left": "18mm", "right": "18mm"},
                print_background=True,
            )
            browser.close()
        html_tmp.unlink(missing_ok=True)
    except Exception as exc:
        print(f"[PDF] Playwright PDF failed ({exc}), keeping HTML file at {html_tmp}")
        output_path = html_tmp  # Return HTML as fallback

    return output_path


def _build_resume_html(packet: ApplicationPacket) -> str:
    """Build a clean, professional HTML resume from the packet."""
    job = packet.job
    score = packet.score

    form_rows = "".join(
        f"<tr><td><strong>{k}</strong></td><td>{v}</td></tr>"
        for k, v in packet.form_answers.items()
    )
    notes_items = "".join(f"<li>{n}</li>" for n in packet.resume_notes)
    matched_terms = ", ".join(score.matched_terms[:12]) if score.matched_terms else "—"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{job.title} — {job.company}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: 'Inter', 'Segoe UI', sans-serif;
    font-size: 11pt;
    color: #1a1a2e;
    line-height: 1.6;
    background: #fff;
    padding: 0;
  }}
  .page {{ padding: 24mm 20mm; max-width: 210mm; margin: auto; }}
  h1 {{ font-size: 22pt; color: #0f3460; border-bottom: 3px solid #0f3460; padding-bottom: 6px; margin-bottom: 4px; }}
  .meta {{ color: #666; font-size: 9.5pt; margin-bottom: 18px; }}
  .score-badge {{
    display: inline-block;
    background: #0f3460;
    color: white;
    border-radius: 4px;
    padding: 2px 10px;
    font-size: 9pt;
    margin-left: 10px;
    vertical-align: middle;
  }}
  h2 {{ font-size: 13pt; color: #0f3460; margin-top: 18px; margin-bottom: 6px;
        border-left: 4px solid #e94560; padding-left: 8px; }}
  p {{ margin-bottom: 8px; }}
  ul {{ padding-left: 18px; margin-bottom: 8px; }}
  li {{ margin-bottom: 3px; }}
  table {{ width: 100%; border-collapse: collapse; margin-top: 6px; font-size: 10pt; }}
  td {{ padding: 5px 8px; border-bottom: 1px solid #eee; vertical-align: top; }}
  td:first-child {{ width: 38%; color: #333; }}
  .tags {{ display: flex; flex-wrap: wrap; gap: 5px; margin-top: 6px; }}
  .tag {{
    background: #e8f0fe;
    color: #0f3460;
    padding: 2px 8px;
    border-radius: 12px;
    font-size: 9pt;
    font-weight: 500;
  }}
  .cover-letter {{ white-space: pre-wrap; background: #f8f9fa; border-radius: 6px;
    padding: 12px 14px; border-left: 3px solid #0f3460; font-size: 10.5pt; }}
  .footer {{ margin-top: 24px; font-size: 8pt; color: #aaa; text-align: center; }}
</style>
</head>
<body>
<div class="page">
  <h1>{job.title} <span class="score-badge">Match {score.match_percent}%</span></h1>
  <div class="meta">
    {job.company} &nbsp;·&nbsp; {job.location}
    {'&nbsp;·&nbsp; 🌐 Remote' if job.remote else ''}
    &nbsp;·&nbsp; <a href="{job.url}">{job.url[:60]}{'…' if len(job.url) > 60 else ''}</a>
  </div>

  <h2>Tailored Summary</h2>
  <p>{packet.tailored_resume_text or score.explanation()}</p>

  <h2>Matched Keywords</h2>
  <div class="tags">
    {''.join(f'<span class="tag">{t}</span>' for t in score.matched_terms[:15]) or '<span>—</span>'}
  </div>

  <h2>Cover Letter</h2>
  <div class="cover-letter">{packet.cover_letter or '—'}</div>

  {f'''<h2>Cold Email</h2>
  <p><strong>Subject:</strong> {packet.cold_email_subject}</p>
  <div class="cover-letter">{packet.cold_email_body}</div>''' if packet.cold_email_subject else ''}

  {f'''<h2>Application Form Answers</h2>
  <table>{form_rows}</table>''' if form_rows else ''}

  {f'''<h2>Resume Notes</h2>
  <ul>{notes_items}</ul>''' if notes_items else ''}

  <div class="footer">Generated by JobFlow · {job.source} · Score {score.score:.3f}</div>
</div>
</body>
</html>"""


class TailorEngine:
    def __init__(self, config: AppConfig, profile: Profile):
        self.config = config
        self.profile = profile
        self.resume_text = extract_resume_text(profile, config.root_dir)

    def build_packet(self, score: JobScore) -> ApplicationPacket:
        questions = extract_form_questions(score.job)
        answers = default_form_answers(self.profile)
        if self.config.llm_dry_run or not self.config.deepseek_api_key:
            return self._build_fallback_packet(score, questions, answers)
        payload = self._generate_with_llm(score.job, questions)
        merged_answers = {**answers, **payload.get("form_answers", {})}
        return ApplicationPacket(
            job=score.job,
            score=score,
            resume_notes=payload.get("resume_notes", []) or self._fallback_notes(score),
            tailored_resume_text=payload.get("tailored_resume_text", ""),
            cover_letter=payload.get("cover_letter", ""),
            cold_email_subject=payload.get("cold_email_subject", ""),
            cold_email_body=payload.get("cold_email_body", ""),
            form_answers=merged_answers,
            form_questions=questions,
        )

    def _build_fallback_packet(
        self,
        score: JobScore,
        questions: list[str],
        answers: dict[str, str],
    ) -> ApplicationPacket:
        role_line = f"Target role alignment: {score.job.title} at {score.job.company}"
        tailored_resume = (
            f"Summary focus: {score.job.title}\n"
            f"Relevant skills: {', '.join(score.matched_terms[:8]) or ', '.join(self.profile.skills[:6])}\n"
            f"Suggested emphasis: automation, APIs, delivery impact, and measurable outcomes."
        )
        cover_letter = (
            f"Dear Hiring Team,\n\n"
            f"I am excited to apply for the {score.job.title} role at {score.job.company}. "
            f"My background aligns well with the role through experience in {', '.join(self.profile.skills[:4])}. "
            f"I am especially interested in this opportunity because it matches my focus on {', '.join(self.profile.target_roles[:2]) or 'software engineering'}.\n\n"
            f"I would bring a practical, execution-focused approach and would welcome the chance to contribute.\n\n"
            f"Sincerely,\n{self.profile.name}"
        )
        cold_email_subject = f"{self.profile.name} for {score.job.title}"
        cold_email_body = (
            f"Hi {score.job.company} team,\n\n"
            f"I found the {score.job.title} role and it looks closely aligned with my background in "
            f"{', '.join(self.profile.skills[:4])}. I am sharing a tailored summary and would love to be considered "
            f"for the role if the opening is still active.\n\n"
            f"Strong fit areas: {', '.join(score.matched_terms[:6]) or ', '.join(self.profile.target_roles[:2])}.\n\n"
            f"Best,\n{self.profile.name}"
        )
        return ApplicationPacket(
            job=score.job,
            score=score,
            resume_notes=self._fallback_notes(score) + [role_line],
            tailored_resume_text=tailored_resume,
            cover_letter=cover_letter,
            cold_email_subject=cold_email_subject,
            cold_email_body=cold_email_body,
            form_answers=answers,
            form_questions=questions,
        )

    def _fallback_notes(self, score: JobScore) -> list[str]:
        notes = []
        if score.matched_terms:
            notes.append("Matched terms: " + ", ".join(score.matched_terms[:8]))
        notes.append("Highlight direct experience with automation, APIs, and shipping local tooling.")
        return notes

    def _generate_with_llm(self, job: JobListing, questions: list[str]) -> dict[str, Any]:
        system = (
            "You create concise, truthful application materials. "
            "Return strict JSON with keys: resume_notes, tailored_resume_text, cover_letter, cold_email_subject, cold_email_body, form_answers."
        )
        prompt = {
            "candidate_profile": {
                "name": self.profile.name,
                "headline": self.profile.headline,
                "location": self.profile.location,
                "target_roles": self.profile.target_roles,
                "skills": self.profile.skills,
                "summary": self.profile.summary,
            },
            "resume_text": self.resume_text[:12000],
            "job": {
                "title": job.title,
                "company": job.company,
                "location": job.location,
                "description": _strip_html(job.description)[:12000],
                "url": job.url,
            },
            "form_questions": questions,
        }
        body = json.dumps(
            {
                "model": self.config.deepseek_model,
                "max_tokens": 1800,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
                ],
                "response_format": {"type": "json_object"},
            }
        ).encode("utf-8")
        request = urllib.request.Request(
            "https://api.deepseek.com/chat/completions",
            data=body,
            method="POST",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.config.deepseek_api_key}",
            },
        )
        with urllib.request.urlopen(request, timeout=60) as response:
            payload = json.loads(response.read().decode("utf-8"))
        choices = payload.get("choices", [])
        text = choices[0].get("message", {}).get("content", "") if choices else ""
        return _parse_json_object(text)


def _strip_html(value: str) -> str:
    value = re.sub(r"<[^>]+>", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _parse_json_object(value: str) -> dict[str, Any]:
    value = value.strip()
    if value.startswith("```"):
        value = value.strip("`")
        if "\n" in value:
            value = value.split("\n", 1)[1]
    start = value.find("{")
    end = value.rfind("}")
    if start == -1 or end == -1:
        return {}
    try:
        parsed = json.loads(value[start : end + 1])
    except json.JSONDecodeError:
        return {}
    return parsed if isinstance(parsed, dict) else {}
